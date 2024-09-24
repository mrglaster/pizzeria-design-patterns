import os

from modules.domain.report.report.base.abstract_report import ComplexReport
from modules.domain.report.report_format.report_format import ReportFormat
from modules.exception.bad_argument_exception import BadArgumentException
from modules.service.managers.settings_manager import SettingsManager
from modules.validation.data_validator import DataValidator
from docx import Document


class ReportDOCX(ComplexReport):
    def __init__(self):
        super().__init__()
        self.format = ReportFormat.FORMAT_DOCX
        self._document = None

    def create(self, data: list):
        DataValidator.validate_field_type(data, list)
        if not len(data):
            raise BadArgumentException("Empty data list provided!")

        used_model = data[0]
        class_name = type(used_model).__name__
        self._document = Document()
        self._document.add_heading(f'{class_name} Report', level=1)
        fields = list(filter(lambda x: not x.startswith("_") and not callable(getattr(used_model.__class__, x)),
                             dir(used_model.__class__)))
        table = self._document.add_table(rows=1, cols=len(fields))
        hdr_cells = table.rows[0].cells
        for idx, field in enumerate(fields):
            hdr_cells[idx].text = field

        for row in data:
            row_cells = table.add_row().cells
            for idx, field in enumerate(fields):
                value = getattr(row, field)
                row_cells[idx].text = str(value)

    def save(self, file_name: str):
        DataValidator.validate_field_type(file_name, str)
        DataValidator.validate_str_not_empty(file_name)
        if self._document is None:
            raise ValueError("No data to save. Please create the report first.")
        save_path = file_name
        if os.path.basename(file_name) == file_name:
            sm = SettingsManager()
            sm.read_settings()
            reps = os.path.join(os.getcwd().replace('test/', ''), sm.settings.reports_path)
            save_path = os.path.join(reps, file_name)
            save_path = save_path.replace('test/', '')
        self._document.save(save_path)
